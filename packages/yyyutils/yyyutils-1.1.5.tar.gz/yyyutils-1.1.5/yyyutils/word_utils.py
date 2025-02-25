from docx import Document
import re
import os
from lxml import etree
from .decorator_utils import DecoratorUtils
from os_utils import OSUtils
import comtypes.client


class WordUtils:
    """
    用于操作word文档的工具类
    """

    def __init__(self, file_path, print_xml_tree=False):
        if not file_path.endswith('.docx'):
            raise ValueError('File must be a.docx file')
        if not os.path.exists(file_path):
            # 新建文档
            self.document = Document()
            self.document.save(file_path)
        self.file_path = file_path
        self.print_xml_tree = print_xml_tree
        self.document = Document(file_path)
        self.xml_body, self.sections, self.paragraphs, self.text_runs, self.text_runs_texts, self.rPrs, self.brs = self.__analysis_to_xml(
            print_xml_tree)
        print(len(self.brs))

    def __analysis_to_xml(self, print_xml_tree=False):
        xml_body = self.document.element.body
        sections = xml_body.findall('.//w:sectPr', namespaces=self.document.element.nsmap)
        paragraphs = xml_body.findall('.//w:p', namespaces=self.document.element.nsmap)
        text_runs = xml_body.findall('.//w:r', namespaces=self.document.element.nsmap)
        text_runs_texts = xml_body.findall('.//w:t', namespaces=self.document.element.nsmap)
        rPrs = xml_body.findall('.//w:rPr', namespaces=self.document.element.nsmap)
        brs = xml_body.findall('.//w:br', namespaces=self.document.element.nsmap)
        print('newTree:\n', etree.tostring(xml_body, encoding='unicode', pretty_print=True)) if print_xml_tree else None
        return xml_body, sections, paragraphs, text_runs, text_runs_texts, rPrs, brs

    def __update_xml(self):
        self.xml_body, self.sections, self.paragraphs, self.text_runs, self.text_runs_texts, self.rPrs, self.brs = self.__analysis_to_xml(
            print_xml_tree=self.print_xml_tree)

    def __create_section(self):
        # 创建新的节元素
        p = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
        pPr = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        sectPr = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr")
        pgSz = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgSz", w="11906", h="16838")
        pgMar = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgMar",
                              top="1440", right="1800", bottom="1440", left="1800", header="851", footer="992",
                              gutter="0")
        cols = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols", space="425", num="1")
        docGrid = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docGrid", type="lines",
                                linePitch="312", charSpace="0")

        # 将这些元素添加到节属性中
        sectPr.append(pgSz)
        sectPr.append(pgMar)
        sectPr.append(cols)
        sectPr.append(docGrid)

        return sectPr

    def __create_run(self, text):
        # 创建新的run元素
        r = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        rFonts = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
        rFonts.set("hint", "eastAsia")
        lang = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lang")
        lang.set("val", "en-US")
        lang.set("eastAsia", "zh-CN")
        rPr = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
        t = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t.text = text

        # 将这些元素添加到run属性中
        rPr.append(rFonts)
        rPr.append(lang)
        r.append(rPr)
        r.append(t)
        return r

    def __create_paragraph(self, section=None, run=None):
        # 创建新的段落元素
        p = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
        pPr = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        rPr = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
        rFonts = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
        rFonts.set("hint", "default")
        lang = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lang")
        lang.set("val", "en-US")
        lang.set("eastAsia", "zh-CN")
        rPr.append(rFonts)
        rPr.append(lang)
        pPr.append(rPr)
        if section is not None:
            pPr.append(section)
            p.append(pPr)
        else:
            p.append(pPr)
        if run is not None:
            p.append(run)
        book_mark_Start = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart")
        book_mark_Start.set("id", "0")
        book_mark_Start.set("name", "_GoBack")
        book_mark_End = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkEnd")
        book_mark_End.set("id", "0")
        p.append(book_mark_Start)
        p.append(book_mark_End)
        return p

    def __create_split_page(self):
        p = self.__create_paragraph()
        run = self.__create_run('')
        br = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br")
        br.set("type", "page")
        run.append(br)
        p.append(run)
        return p

    def insert_text_to_section(self, text, section=None, style=None, change_line=False):
        """
        在指定节中插入文本，如果section参数为None，则插入到最后一个节中，如果section参数为int，则插入到第section个节中
        :param text:
        :param section:
        :param style:
        :param change_line:
        :return:
        """
        if section is None or section >= len(self.sections):
            if change_line:
                # 最后一个段落的后面添加一个新的段落，并在新的段落中插入run
                self.paragraphs[-1].addnext(self.__create_paragraph(run=self.__create_run(text)))
            else:
                self.paragraphs[-1].append(self.__create_run(text))
        elif isinstance(section, int):
            s = self.sections[section - 1]

            p = s.getparent().getparent()
            s_t = p.find('.//w:t', namespaces=self.document.element.nsmap)
            s_text = s_t.text
            if change_line:
                new_p = self.__create_paragraph(run=self.__create_run(s_text))
                p.addprevious(new_p)
                s_t.text = text
            else:
                s_t.text = s_text + text
        else:
            print('section参数错误')

        self.document.save(self.file_path)
        self.__update_xml()

    def delete_single_section(self, section=-1):
        if len(self.sections) == 1:
            raise ValueError('文档中只有一个节，不能删除')
        if section >= len(self.sections) or section < 0:
            del_section = self.sections[-2]
            del_section_p = del_section.getparent().getparent()
            del_section_t = del_section_p.find('.//w:t', namespaces=self.document.element.nsmap)
            new_p = self.__create_paragraph(run=self.__create_run(del_section_t.text))
            # 删除此时self.p中在del_section_p之后的所有p标签
            for p in self.paragraphs[self.paragraphs.index(del_section_p) + 1:]:
                p.getparent().remove(p)
            # 添加新的p标签，删除del_section_p
            del_section_p.addnext(new_p)
            del_section_p.getparent().remove(del_section_p)
        else:
            del_section = self.sections[section - 1]
            del_section_p = del_section.getparent().getparent()

            if section == 1:
                for p in self.paragraphs[0:self.paragraphs.index(del_section_p) + 1]:
                    p.getparent().remove(p)
            else:
                pre_section = self.sections[section - 2]
                pre_section_p = pre_section.getparent().getparent()
                for p in self.paragraphs[
                         self.paragraphs.index(pre_section_p) + 1:self.paragraphs.index(del_section_p) + 1]:
                    p.getparent().remove(p)
        self.document.save(self.file_path)
        self.__update_xml()

    def insert_single_section(self, insert_position=-1):
        """
        新建一个节，相当于把原有节里面的内容刷到新的节里，然后再原有节和新节之间插入一个空行
        :param insert_position:
        :return:
        """
        p = self.__create_paragraph(self.__create_section())
        if insert_position >= len(self.sections) or insert_position < 0:
            # 此时，需要计算出最后一个节和前一个节之间p标签的个数，然后在倒数第二个p标签之后插入一个新的节，并把最后一个p标签的文本存入这个节的run里面，之后删除最后一个p标签
            del_p = self.paragraphs[-1]
            pre_text = del_p.text
            del_p.getparent().remove(del_p)
            self.xml_body.insert(len(self.sections) - 1,
                                 self.__create_paragraph(self.__create_section(), self.__create_run(pre_text)))
            self.__update_xml()
            self.xml_body.insert(len(self.sections) - 1, self.__create_paragraph())
        else:
            self.xml_body.insert(insert_position, p)

        # 保存文档
        self.document.save(self.file_path)
        self.__update_xml()

    def __add_style(self, rPr, *styles):
        font_name, font_size, font_color, bold, italic, underline = styles
        rFonts = rPr.find('.//w:rFonts', namespaces=self.document.element.nsmap)
        if font_name is not None:
            if rFonts is not None:
                # 删除rFonts标签，重新写一个
                rPr.remove(rFonts)
            rFonts = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
            rFonts.set("hint", "eastAsia")
            rFonts.set("ascii", font_name)
            rFonts.set("hAnsi", font_name)
            rFonts.set("eastAsia", font_name)
            rFonts.set("cs", font_name)
            rPr.append(rFonts)
        if font_size is not None:
            for tag in ['sz', 'szCs']:
                element = rPr.find(f'.//w:{tag}', namespaces=self.document.element.nsmap)
                if element is not None:
                    # 删除sz标签，重新写一个
                    rPr.remove(element)
                element = etree.Element(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{tag}")
                element.set("val", str(font_size * 2))
                rPr.append(element)
        if font_color is not None:
            for tag in ['color', 'colorCs']:
                element = rPr.find(f'.//w:{tag}', namespaces=self.document.element.nsmap)
                if element is not None:
                    # 删除color标签，重新写一个
                    rPr.remove(element)
                element = etree.Element(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{tag}")
                element.set("val", f"#{font_color[0]:02x}{font_color[1]:02x}{font_color[2]:02x}")
                rPr.append(element)
        if bold:
            if rPr.find('.//w:b', namespaces=self.document.element.nsmap) is None:
                b = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b")
                rPr.append(b)
        else:
            b = rPr.find('.//w:b', namespaces=self.document.element.nsmap)
            if b is not None:
                rPr.remove(b)
        if italic:
            if rPr.find('.//w:i', namespaces=self.document.element.nsmap) is None:
                i = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}i")
                rPr.append(i)
        else:
            i = rPr.find('.//w:i', namespaces=self.document.element.nsmap)
            if i is not None:
                rPr.remove(i)
        if underline:
            if rPr.find('.//w:u', namespaces=self.document.element.nsmap) is None:
                u = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u")
                u.set('val', 'single')
                rPr.append(u)
        else:
            u = rPr.find('.//w:u', namespaces=self.document.element.nsmap)
            if u is not None:
                rPr.remove(u)

    @DecoratorUtils.validate_input
    def set_text_style(self, text=None, font_name=None, font_size: int = None, font_color=None, bold=False,
                       italic=False,
                       underline=False):
        """
        设置文本的字体，字号，颜色，粗体，斜体，下划线样式
        :param text:
        :param style_name:
        :return:
        """
        if text is None:
            for rPr in self.rPrs:
                self.__add_style(rPr, font_name, font_size, font_color, bold, italic, underline)
        else:
            pass

        self.document.save(self.file_path)
        self.__update_xml()

    def insert_single_page(self, insert_position=-1):
        p = self.__create_split_page()
        if insert_position > len(self.brs) or insert_position < 0:
            print(123)
            self.sections[-1].addprevious(p)
            self.__update_xml()
            self.sections[-1].addprevious(self.__create_paragraph())
        else:
            self.brs[insert_position - 1].getparent().getparent().addnext(p)

        self.document.save(self.file_path)
        self.__update_xml()

    def insert_text_to_page(self, text, page_num):
        pass

    def add_custom_style(self, document, style):
        try:
            style.base_style = document.styles['Normal Table']
        except KeyError:
            print("Style 'Normal Table' not found. Using default style.")
            style.base_style = document.styles['Normal']

        # 设置表格样式属性
        style.paragraph_format.alignment = 1  # 居中对齐
        style.font.name = 'Arial'
        style.font.size = 120000  # 12pt

    def add_single_table(self, rows=2, cols=3):
        table = self.document.add_table(rows=rows, cols=cols)
        style = table.style
        self.add_custom_style(self.document, style)
        self.document.save(self.file_path)

    def set_section_footer(self, section=None, text=None):
        section = self.document.sections[-1]  # 获取文档的最后一个节，即新创建的节
        footer = section.footer
        footer.is_linked_to_previous = False  # 断开与前一节的链接
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = text
        self.document.save(self.file_path)

    def copy_docx(self, source_file_path):
        """
        复制word文档,如果目标文件存在,则覆盖
        :return:
        """
        if os.path.exists(self.file_path):
            os.remove(self.file_path)
        document = Document(source_file_path)
        document.save(self.file_path)

    def regex_replace(self, pattern, replace_str):
        """
        使用正则表达式替换文档中的文本
        :param pattern: 正则表达式
        :param replace_str: 替换字符串
        :return:
        """
        for paragraph in self.document.paragraphs:
            paragraph.text = re.sub(pattern, replace_str, paragraph.text)
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = re.sub(pattern, replace_str, paragraph.text)
        self.document.save(self.file_path)

    @staticmethod
    def convert_to_pdf(files_path, outputs_path_folder, need_text=False):
        """
        将word文档转换为pdf文档
        :param files_path: 输入文件或文件列表的路径
        :param outputs_path_folder: 输出文件夹的路径
        :param need_text: 是否需要保留备注文字
        :return: 成功和失败的文件列表
        """
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        success_files = []
        failed_files = []

        if isinstance(files_path, str):
            files_path = [files_path]

        if not os.path.exists(outputs_path_folder):
            os.makedirs(outputs_path_folder)

        outputs_path = [os.path.join(outputs_path_folder, os.path.splitext(os.path.basename(file_path))[0] + '.pdf')
                        for file_path in files_path]

        for file_path, output_path in zip(files_path, outputs_path):
            try:
                if not os.path.exists(file_path):
                    print(f"文件不存在: {file_path}")
                    failed_files.append(file_path)
                    continue
                doc = word.Documents.Open(file_path)
                if not need_text:
                    doc.Content.Delete()
                output_path = OSUtils.autochange_file_name_while_exist(outputs_path_folder, output_path)
                doc.SaveAs(output_path, FileFormat=17)
                doc.Close()
                if os.path.exists(output_path):
                    print(f"{file_path}转换成功")
                    success_files.append(file_path)
                else:
                    print(f"{file_path}转换失败")
                    failed_files.append(file_path)
            except Exception as e:
                print(f"{file_path}转换时发生错误: {str(e)}")
                failed_files.append(file_path)
        word.Quit()
        return success_files, failed_files


if __name__ == '__main__':

    paths, _, names = OSUtils.get_all_files_in_directory_by_name_extension(r'D:\学生工作\2024暑期三下乡\乡村振兴',
                                                                           'docx')
    print(names)
    paths1 = OSUtils.get_all_files_in_directory_by_name_or_name_prefix(r'D:\学生工作\2024暑期三下乡\乡村振兴', '~$')[0]
    for path in paths1:
        OSUtils.delete_single_file(path)
    output_paths = r'C:\Desktop\output'
    WordUtils.convert_to_pdf(paths, output_paths)
